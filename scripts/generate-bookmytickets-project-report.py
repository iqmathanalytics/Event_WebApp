"""
Generate Book My Tickets project report (Word) following the sample SIH report format.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_document_defaults(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    p.paragraph_format.space_after = Pt(4)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    p.paragraph_format.space_after = Pt(14)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(6)


def add_body(doc, text):
    doc.add_paragraph(text)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.left_indent = Inches(0.25)


def add_label_block(doc, label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    p.add_run(value)


def add_sdg_section(doc, goal, targets):
    add_heading(doc, goal, level=2)
    for t in targets:
        add_body(doc, t)


def build_report():
    doc = Document()
    set_document_defaults(doc)

    # Cover / header
    add_title(doc, "Book My Tickets — City Events & Lifestyle Hub")
    add_subtitle(doc, "Full-Stack Web Platform for Event Discovery, On-Site Ticketing, and Organizer Operations")
    add_subtitle(doc, "Project Report | bookmytickets.us")

    add_heading(doc, "Abstract")
    add_body(
        doc,
        "Book My Tickets is a production-grade web platform that helps residents discover events, deals, "
        "and local creators in their city while enabling organizers to publish listings, sell tickets on-site, "
        "manage reserved seating, and track performance. The solution combines a React single-page application "
        "with a Node.js REST API, TiDB Cloud database, Stripe payments, Seats.io seating charts, Brevo "
        "transactional email, Google Analytics insights, and Cloudinary media hosting. The platform supports "
        "guest and registered checkout, coupon-based promotions, QR-based venue check-in, and a full admin "
        "moderation workflow—delivering an end-to-end digital ecosystem for city-scale event commerce."
    )

    add_heading(doc, "Introduction")
    add_body(
        doc,
        "Urban audiences increasingly rely on digital channels to find concerts, festivals, workshops, "
        "and community gatherings. However, discovery is often fragmented across social media, ticketing "
        "aggregators, and organizer websites. Small and mid-size event organizers also lack affordable tools "
        "to list events, accept payments, assign seats, issue digital tickets, and understand audience behavior."
    )
    add_body(
        doc,
        "Book My Tickets addresses this gap by unifying event discovery, lifestyle content (deals and "
        "influencers), and on-platform ticket sales in one moderated marketplace. Users browse by city, "
        "compare ticket tiers, apply coupons, and complete checkout securely. Organizers manage events, "
        "design seating charts, export bookings, and view analytics. Administrators maintain content quality, "
        "user access, newsletter communications, and ticket verification at the venue."
    )

    add_heading(doc, "Problem Statement")
    add_body(
        doc,
        "City residents struggle to find trustworthy, up-to-date local events in one place. Organizers face "
        "high friction when moving from event promotion to payment collection, especially for multi-tier and "
        "reserved-seat events. Existing generic listing sites rarely offer integrated checkout, seat holds, "
        "coupon logic, check-in QR codes, and role-based dashboards in a single product tailored to US city markets."
    )

    add_heading(doc, "Need for the Solution")
    add_body(doc, "There is a need for a unified platform that can:")
    add_bullets(
        doc,
        [
            "Aggregate city-filtered events, deals, and creator profiles in one trusted interface",
            "Enable organizers to submit, edit, and publish events through an approval workflow",
            "Support both external ticket links and native on-site checkout with Stripe",
            "Handle reserved seating with real-time seat selection and hold management",
            "Apply percentage and fixed-amount coupon codes with usage limits and expiry rules",
            "Deliver booking confirmations, QR codes, and guest-account onboarding via email",
            "Provide admin moderation, booking exports, newsletter tools, and QR ticket verification",
            "Offer organizer insights using Google Analytics Data API for traffic and engagement trends",
        ],
    )

    add_heading(doc, "Proposed Solution")
    add_body(doc, "The Book My Tickets platform provides:")
    add_bullets(
        doc,
        [
            "Public discovery portal with city selector, search, filters, favorites, and featured listings",
            "Rich event pages with venue maps, galleries, promo videos, and ticket tier breakdown",
            "On-site checkout for general admission and reserved seating events",
            "Stripe Payment Intents with guest checkout and signed-in user flows",
            "Seats.io integration for chart design, channel management, seat holds, and booking",
            "Organizer workspace for event CRUD, coupon management, booking export, and insights",
            "Admin console for moderation, user/team management, cities, communications, and check-in",
            "Transactional email via Brevo and optional Mailchimp newsletter synchronization",
            "Cloudinary-backed image uploads and SEO-friendly public listing slugs",
        ],
    )

    add_heading(doc, "Technical Architecture")
    add_label_block(
        doc,
        "Frontend",
        "React 18, Vite, Tailwind CSS, React Router, Axios, Context API, Framer Motion, Recharts, Swiper",
    )
    add_label_block(doc, "Backend", "Node.js, Express.js, Zod validation, JWT auth, Helmet, CORS, rate limiting")
    add_label_block(doc, "Database", "TiDB Cloud (MySQL-compatible) with SQL migrations")
    add_label_block(doc, "Payments", "Stripe (Payment Intents, webhooks, wallet domain support)")
    add_label_block(doc, "Seating", "Seats.io (designer, event manager, server-side holds)")
    add_label_block(doc, "Email", "Brevo transactional email; optional Mailchimp audience sync")
    add_label_block(doc, "Analytics", "Google Analytics 4 (client) + GA4 Data API (organizer insights)")
    add_label_block(doc, "Media", "Cloudinary image hosting and upload API")
    add_label_block(
        doc,
        "Deployment",
        "Production on MilesWeb (Node API + static frontend); alternative Render API + Netlify SPA",
    )

    add_heading(doc, "Development Process")
    add_numbered(
        doc,
        [
            "Requirement analysis for public users, organizers, and administrators",
            "Database schema design and SQL migration setup (events, bookings, coupons, cities)",
            "REST API development with validation, auth middleware, and service layer",
            "React frontend with role-based dashboards and responsive UI",
            "Stripe checkout and webhook fulfillment integration",
            "Seats.io seating designer, buyer chart, holds, and tier mapping logic",
            "Coupon engine with holds, redemption limits, and checkout pricing",
            "Email templates, QR check-in, and admin verification tools",
            "Production deployment preparation, environment verification, and go-live checklist",
        ],
    )

    add_heading(doc, "Key Features")
    add_bullets(
        doc,
        [
            "City-based event, deal, and influencer discovery",
            "Single-day and multi-day event scheduling with ticket levels (General, Premium, VIP)",
            "External ticket links or platform-native checkout",
            "Reserved seating with interactive seat map and timed seat holds",
            "Coupon codes: percentage/fixed discount, scope, caps, and per-user limits",
            "Guest checkout with optional account creation and booking email CTA",
            "User dashboard: profile, favorites, bookings, and submissions",
            "Organizer dashboard: events, coupons, bookings export, seating tools, insights",
            "Admin dashboard: moderation, analytics, users, newsletter, contact, ticket scanner",
            "Favorites, newsletter subscription, and contact messaging",
            "Google Sign-In and JWT session management with refresh tokens",
            "Public QR ticket images and admin check-in verification",
        ],
    )

    add_heading(doc, "System Modules")
    add_label_block(doc, "Public modules", "Home, Events, Deals, Influencers, Contact, Newsletter, Auth")
    add_label_block(
        doc,
        "Organizer modules",
        "Event form, ticket levels, seating designer/channels, coupons, bookings, insights",
    )
    add_label_block(
        doc,
        "Admin modules",
        "Moderation queue, listings, users/team, bookings, communications, cities, verify-ticket",
    )
    add_label_block(
        doc,
        "Core API areas",
        "/auth, /users, /events, /bookings, /deals, /influencers, /favorites, /newsletter, /admin, /meta, /webhooks/stripe",
    )

    add_heading(doc, "Database Design (Major Entities)")
    add_bullets(
        doc,
        [
            "users, user_onboarding_profiles — accounts, roles, and capability flags",
            "cities, categories — location and taxonomy metadata",
            "events — listings, schedules, ticket config, seating mode, moderation status",
            "event_bookings, event_checkout_payments — orders, payments, QR check-in",
            "event_coupons, event_coupon_holds, event_coupon_redemptions — promotion engine",
            "deals, influencers, dealer_profiles, services — lifestyle marketplace content",
            "favorites, newsletter_subscribers, contact_messages — engagement and comms",
            "admin_notifications, platform_ticket_access_requests — operations workflow",
        ],
    )

    add_heading(doc, "Security & Authentication")
    add_bullets(
        doc,
        [
            "JWT access and refresh tokens with secure password hashing (bcrypt)",
            "Role-based access: user, organizer, admin with fine-grained capability flags",
            "Separate user and staff login flows; Google OAuth with server-side token verification",
            "Zod request validation on API inputs; Helmet security headers and CORS allowlist",
            "Stripe webhook signature verification; rate limiting on sensitive public endpoints",
            "Trust proxy configuration for correct client IP behind production reverse proxy",
        ],
    )

    add_heading(doc, "Expected Impact")
    add_body(
        doc,
        "The platform strengthens local event economies by giving organizers a professional digital storefront "
        "and giving residents a single trusted guide to city life. Integrated checkout reduces abandonment "
        "compared with redirecting users to external sites. Reserved seating and coupon tools improve revenue "
        "management for venues. Admin moderation and QR check-in increase operational trust. Analytics help "
        "organizers understand demand patterns and optimize marketing spend. Overall, Book My Tickets lowers "
        "the barrier for community events to go digital while improving the attendee experience."
    )

    add_heading(doc, "Future Scope")
    add_bullets(
        doc,
        [
            "Full public launch of the Services module (backend and page already partially built)",
            "Mobile-optimized PWA or native apps for ticket wallet and push notifications",
            "Deeper CRM for organizers: attendee segmentation, email campaigns, and repeat-buyer offers",
            "Multi-language and multi-currency support for broader US metro expansion",
            "Waitlists, transfers, and partial refunds for high-demand seated events",
            "Partner APIs for venue POS and third-party promoter integrations",
            "Automated CI/CD, test coverage, and performance monitoring dashboards",
            "AI-assisted event tagging, pricing recommendations, and demand forecasting",
        ],
    )

    add_heading(doc, "Conclusion")
    add_body(
        doc,
        "Book My Tickets delivers a complete digital stack for city-scale event discovery and commerce. "
        "By combining moderated listings, native payments, reserved seating, promotional coupons, transactional "
        "email, and role-specific dashboards, the platform serves attendees, organizers, and administrators in "
        "one cohesive system. Built on modern web technologies and deployed for production at bookmytickets.us, "
        "the solution is scalable, secure, and extensible—providing a strong foundation for growing local event "
        "ecosystems and lifestyle engagement across US cities."
    )

    add_heading(doc, "Sustainable Development Goals (SDGs)")
    add_sdg_section(
        doc,
        "Primary Goal: SDG 11 – Sustainable Cities and Communities",
        [
            "Target 11.3 (Inclusive and Sustainable Urbanization): City-filtered discovery helps residents "
            "participate in local cultural and community events, strengthening urban social fabric.",
            "Target 11.6 (Reduce Environmental Impact of Cities): Digital ticketing and QR check-in reduce "
            "paper waste and streamline venue entry compared with manual gate processes.",
        ],
    )
    add_sdg_section(
        doc,
        "Secondary Goal: SDG 8 – Decent Work and Economic Growth",
        [
            "Target 8.3 (Promote Productive Activities and Entrepreneurship): Organizers—especially small "
            "venues and independent promoters—gain affordable tools to sell tickets and grow audiences.",
            "Target 8.2 (Achieve Higher Economic Productivity Through Innovation): Integrated analytics, "
            "coupons, and seating management improve operational efficiency and revenue capture.",
        ],
    )
    add_sdg_section(
        doc,
        "Secondary Goal: SDG 9 – Industry, Innovation and Infrastructure",
        [
            "Target 9.5 (Enhance Scientific Research and Technological Capabilities): The platform "
            "demonstrates modern full-stack engineering with payments, real-time seating, and data APIs.",
            "Target 9.c (Increase Access to ICT): A responsive web application broadens access to event "
            "information and ticketing for users across devices and connectivity levels.",
        ],
    )
    add_sdg_section(
        doc,
        "Secondary Goal: SDG 10 – Reduced Inequalities",
        [
            "Target 10.2 (Social, Economic and Political Inclusion): Guest checkout and clear pricing "
            "with coupon support improve access to events for users without prior platform accounts.",
        ],
    )

    add_heading(doc, "Impact Statement")
    add_body(
        doc,
        "Book My Tickets supports inclusive, innovation-driven local economies by connecting communities "
        "with experiences in their city while empowering organizers with enterprise-grade ticketing tools. "
        "Through secure digital payments, transparent moderation, and data-informed organizer insights, the "
        "platform contributes to more vibrant, accessible, and sustainable urban entertainment ecosystems."
    )

    add_heading(doc, "Prototype / Application Screens")
    add_bullets(
        doc,
        [
            "Landing & Home Page: Hero slideshow, city filter, featured events, deals, and influencers",
            "Events Browse & Detail: Filters, maps, galleries, ticket tiers, favorites, and checkout panel",
            "Checkout Flow: Date selection, seat chart (reserved), coupons, Stripe payment, confirmation",
            "User Dashboard: Profile, bookings history, favorites, and submission management",
            "Organizer Dashboard: Event editor, seating designer/channels, coupons, bookings export, insights",
            "Admin Dashboard: Moderation, analytics, users, newsletter, contact inbox, QR ticket scanner",
        ],
    )

    add_heading(doc, "Project Information")
    add_label_block(doc, "Project Name", "Book My Tickets (US Event Website)")
    add_label_block(doc, "Live URL", "https://bookmytickets.us")
    add_label_block(doc, "Repository", "https://github.com/iqmathanalytics/Event_WebApp")
    add_label_block(doc, "Technology Category", "Full-Stack Web Application (MERN-style with TiDB Cloud)")
    add_label_block(doc, "Target Users", "Event attendees, organizers, venue operators, platform administrators")

    return doc


def main():
    out_dir = Path(r"e:\US-Event-Website\docs")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "Book_My_Tickets_Project_Report.docx"
    doc = build_report()
    doc.save(out_path)
    print(f"Created: {out_path}")


if __name__ == "__main__":
    main()
